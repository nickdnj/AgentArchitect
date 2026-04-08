#!/usr/bin/env python3
"""Generate CAGE MATCH Business Plan as a formatted Word document with embedded images."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE = os.path.dirname(os.path.abspath(__file__))

def img(name):
    return os.path.join(BASE, name)

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=10, color=None, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1a1a3e"):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        set_cell_text(cell, h, bold=True, size=9, color=(255, 255, 255))

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f4f4f8")
            bold = isinstance(val, str) and val.startswith("**")
            text = val.replace("**", "") if bold else val
            set_cell_text(cell, text, bold=bold, size=9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3e)
    return h

def add_body(doc, text, bold=False, italic=False, size=11):
    """Add body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Add bullet point, optionally with bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    else:
        p.text = ""
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

def add_image_with_caption(doc, path, caption, width=Inches(6.0)):
    """Add centered image with italic caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(path):
        run.add_picture(path, width=width)
    else:
        run.add_text(f"[Image not found: {path}]")

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Calibri"

def add_callout_box(doc, text, bold_prefix=None):
    """Add a highlighted callout paragraph."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3)
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)

    # Add left border via XML
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="18" w:space="8" w:color="4a6cf7"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)

    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"


# ============================================================
# BUILD DOCUMENT
# ============================================================

doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Narrow margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# TITLE PAGE
# ============================================================

# Add some spacing before title
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CAGE MATCH")
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3e)
run.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Gamified Batting Cage Experience")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4a, 0x6c, 0xf7)
run.font.name = "Calibri"

doc.add_paragraph()

divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run("━" * 40)
run.font.color.rgb = RGBColor(0xa0, 0xc4, 0xff)
run.font.size = Pt(14)

doc.add_paragraph()

for_line = doc.add_paragraph()
for_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = for_line.add_run("Business Plan for")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

batter_up = doc.add_paragraph()
batter_up.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = batter_up.add_run("Batter Up — Bethpage, Long Island")
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3e)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("April 2026  •  Version 1.0  •  Draft for Internal Review")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

prepared = doc.add_paragraph()
prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prepared.add_run("Prepared by Nick DeMarco")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Page break
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================

add_heading_styled(doc, "Table of Contents", level=1)
toc_items = [
    "1.  Executive Summary",
    "2.  The Opportunity",
    "3.  Game Concept — CAGE MATCH",
    "4.  Hardware Architecture",
    "5.  MVP vs. Full Build",
    "6.  Bill of Materials",
    "7.  Financial Projections",
    "8.  Implementation Timeline",
    "9.  Competitive Moat",
    "10. Risks & Mitigations",
    "11. Phase 2 Expansion",
    "12. Visual Mockups",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================

add_heading_styled(doc, "1. Executive Summary", level=1)

add_body(doc, "Cage Match transforms one of Batter Up's existing batting cages into an interactive, gamified experience — a cross between a pinball machine and TopGolf, but for baseball.")

add_body(doc, "Using computer vision, LED targets, and a real-time scoring display, batters compete against themselves and each other in a game that turns 10 pitches into an arcade experience. Spectators watch the action unfold on a display outside the cage. A persistent leaderboard drives repeat visits and social sharing.")

add_callout_box(doc, "Long Island has zero gamified batting cage experiences. Every cage on the island operates with 1980s technology — token in, hit, leave. No scores, no competition, no reason to come back tomorrow. Cage Match changes that.", bold_prefix="The key insight: ")

add_styled_table(doc, ["Metric", "Detail"], [
    ["**MVP Investment**", "~$1,500 (tablet display)"],
    ["**Full Build**", "~$3,000 (outdoor TV)"],
    ["**Proprietary Hardware**", "Custom-built, NVIDIA Jetson + computer vision"],
    ["**Licensing Fees**", "None. Owned outright (HitTrax = $10K–$20K, ProBatter = $30K)"],
    ["**Disruption to Operations**", "Zero — overlays on current token system"],
], col_widths=[2.5, 4.0])

doc.add_paragraph()

# Hero image
add_image_with_caption(doc, img("docmockup_exterior.png"),
    "Concept rendering: The Cage Match cage at dusk. LED targets glow inside while a tablet display shows scores to spectators.")

doc.add_page_break()

# ============================================================
# 2. THE OPPORTUNITY
# ============================================================

add_heading_styled(doc, "2. The Opportunity", level=1)

add_heading_styled(doc, "The Market Gap", level=2)

add_styled_table(doc, ["What Exists on Long Island", "What Does NOT Exist"], [
    ["15+ batting cage facilities", "Zero gamified batting experiences"],
    ["1 HitTrax facility (Bohemia, training-focused)", "No recreational/social hitting venue"],
    ["Token-operated cages unchanged since the 1980s", "No leaderboards, no group competition"],
    ["Baseball culture deeply embedded in LI families", 'No "TopGolf of baseball" anywhere in the NY metro'],
], col_widths=[3.2, 3.2])

doc.add_paragraph()

add_heading_styled(doc, "Proof Points", level=2)
add_bullet(doc, " (Dallas, TX) opened Oct 2025 as the 'TopGolf of baseball' — charges $18–24/person/hour, raised $7.3M, plans 25+ US locations by 2030. No East Coast locations yet.", bold_prefix="Batbox")
add_bullet(doc, " proved that gamifying a commoditized range experience creates a 5–10x revenue premium. There is no TopGolf on Long Island.", bold_prefix="TopGolf")
add_bullet(doc, " cages command a 50–75% pricing premium over standard cages nationally.", bold_prefix="HitTrax")
add_bullet(doc, " transformed bowling from $5/game to $45–65/hour per lane using the same playbook: gamification + time-based pricing + social experience.", bold_prefix="Bowlero")

add_heading_styled(doc, "Historical Inspiration", level=2)
add_body(doc, "Baseball and gaming have been intertwined since the 1920s. The first coin-operated baseball pinball machines appeared in 1929, and by the 1940s, companies like Williams, Gottlieb, and Chicago Coin created an entire genre of pitch-and-bat arcade games. These games succeeded because they combined real skill with arcade scoring, visual feedback, and social competition — the exact elements that make TopGolf work today. Cage Match brings this full circle: a real batting cage with arcade game mechanics.")

add_heading_styled(doc, "Why Batter Up Is Perfectly Positioned", level=2)
add_bullet(doc, " — operating in Bethpage since the 1980s", bold_prefix="Existing customer base and brand recognition")
add_bullet(doc, " — 9 hardball + 4 softball cages, power, lighting", bold_prefix="Physical infrastructure in place")
add_bullet(doc, " — picnic tables, shade sails, social gathering space behind cages", bold_prefix="Spectator area with seating")
add_bullet(doc, " — built-in foot traffic of the target demographic", bold_prefix="Mini golf draws families")
add_bullet(doc, " — can move fast, test, iterate without corporate approval chains", bold_prefix="Family-owned")

# Actual facility photos
doc.add_paragraph()
add_heading_styled(doc, "Batter Up Today", level=2)

add_image_with_caption(doc, img("../images/bu3.jpg"),
    "Batter Up facility — side-by-side chain-link cages with red token machines at each entrance.", width=Inches(5.5))

add_image_with_caption(doc, img("../images/bu8.jpg"),
    "Evening at Batter Up — families line up at the cages under flood lights.", width=Inches(5.5))

doc.add_page_break()

# ============================================================
# 3. GAME CONCEPT
# ============================================================

add_heading_styled(doc, "3. Game Concept — CAGE MATCH", level=1)

add_heading_styled(doc, "How It Works", level=2)

steps = [
    ("1. Walk up. ", "The Cage Match cage looks different — LED targets glow inside. A display outside shows the leaderboard and 'DROP A TOKEN TO PLAY.'"),
    ("2. Drop a token. ", "The token signal triggers the game system. The display shows 'GAME ON' and the pitching machine activates (same as today)."),
    ("3. Targets light up. ", "5 LED target panels on the far end illuminate. Each target is a different color and point value."),
    ("4. Swing. ", "The high-speed camera tracks the ball after contact. Hit a target zone → target flashes, pinball sound plays, points appear on display."),
    ("5. Multiplier builds. ", "Hit 2 targets in a row = 2x. Three = 3x. Max 5x. Miss resets it. This creates the 'one more try' streak mechanic."),
    ("6. 10 pitches, game over. ", "Final score displays with a celebration animation. If the score makes the leaderboard, it goes up on the board."),
    ("7. Leaderboard persists. ", "Daily and all-time high scores visible to everyone walking by."),
]
for prefix, text in steps:
    add_bullet(doc, text, bold_prefix=prefix)

add_heading_styled(doc, "Scoring System", level=2)

add_styled_table(doc, ["Target Zone", "Location", "Points", "Color"], [
    ["Bullseye", "Center", "1,000", "Gold (flashing)"],
    ["Upper Left", "High & pulled", "500", "Red"],
    ["Upper Right", "High & opposite field", "300", "Blue"],
    ["Lower Left", "Line drive pull side", "200", "Green"],
    ["Lower Right", "Line drive opposite", "100", "Yellow"],
], col_widths=[1.5, 2.0, 1.0, 1.5])

doc.add_paragraph()

add_heading_styled(doc, "Bonus Mechanics", level=2)
add_bullet(doc, " 2x after 2 consecutive hits, 3x after 3, max 5x", bold_prefix="Streak Multiplier:")
add_bullet(doc, " Randomly, one target flashes gold for one pitch — worth 2,000 points", bold_prefix="Bonus Target:")
add_bullet(doc, " Hit all 10 pitches into target zones = 5,000 point bonus", bold_prefix="Perfect Game:")
add_bullet(doc, " ~35,000 points (all bonus targets + max multiplier + perfect game)", bold_prefix="Maximum theoretical score:")

# Batter POV image
doc.add_paragraph()
add_image_with_caption(doc, img("docmockup_batter_pov.png"),
    "Batter's view: 5 illuminated LED target panels at the far end of the cage. The gold center target (1,000 pts) is the bullseye.")

doc.add_page_break()

add_heading_styled(doc, "Game Modes", level=2)
add_styled_table(doc, ["Mode", "Phase", "Description"], [
    ["**Target Blitz**", "MVP", "10 pitches, score as many points as possible. Leaderboard competition. 1 token."],
    ["**Derby**", "Phase 2", "2–4 players take turns. Head-to-head. Last place gets 5-pitch redemption at 3x. Birthday/group mode."],
    ["**Innings**", "Phase 2", "9 innings, 3 strikes = out. Zones map to singles/doubles/triples/HRs. Virtual base runners on screen."],
], col_widths=[1.5, 1.0, 4.0])

doc.add_paragraph()

add_heading_styled(doc, "The Social Loop", level=2)
add_callout_box(doc, "\nSee leaderboard → Want to try → Play → Get score →\n    → 'I can beat that' → Play again (more tokens)\n    → Tell friend → Friend comes → New player\n    → Post on social → Organic marketing\n\nThe leaderboard is the engine. It converts every play into either a repeat play or a referral.")

# Tablet UI image
doc.add_paragraph()
add_image_with_caption(doc, img("docmockup_tablet_ui.png"),
    "The Cage Match web app — score, multiplier, pitch counter, and daily leaderboard.", width=Inches(3.5))

doc.add_page_break()

# ============================================================
# 4. HARDWARE ARCHITECTURE
# ============================================================

add_heading_styled(doc, "4. Hardware Architecture", level=1)

add_styled_table(doc, ["Component", "Hardware", "Role"], [
    ["**Brain**", "NVIDIA Jetson Orin Nano Super ($249)", "Computer vision, game logic, web server"],
    ["**Eyes**", "Arducam IMX477 @ 1080p/120fps ($65)", "Ball trajectory tracking via MIPI CSI"],
    ["**Targets**", "ESP32 + WS2812B LEDs + polycarbonate ($150)", "5 illuminated target zones inside cage"],
    ["**Display**", "Weatherproof tablet (MVP) / outdoor TV", "Game state, scores, leaderboard"],
    ["**Trigger**", "Optocoupler on token machine ($8)", "Auto-starts game on token drop"],
    ["**Audio**", "Weatherproof speakers ($85)", "Pinball sounds, crowd reactions"],
    ["**Network**", "4G router or existing WiFi ($89)", "Leaderboard sync, remote monitoring"],
], col_widths=[1.2, 2.8, 2.5])

doc.add_paragraph()

add_heading_styled(doc, "How Ball Tracking Works", level=2)
tracking_steps = [
    "Camera runs at 120 frames per second, mounted behind and above the batter",
    "System watches for bat-ball contact (sudden trajectory change)",
    "OpenCV background subtraction + centroid tracking follows ball across frames",
    "Ball trajectory mapped to 5 target zones (pixel regions in camera FOV)",
    "Zone hit → MQTT message to ESP32 → LED target flashes + sound effect",
    "At 80mph, ball moves ~1.1 ft/frame at 120fps — captured in 3+ frames across cage",
]
for i, step in enumerate(tracking_steps, 1):
    add_bullet(doc, f" {step}", bold_prefix=f"{i}.")

add_heading_styled(doc, "Camera Placement", level=2)
add_body(doc, "Mounted behind and above the batter — the safest location since all balls travel AWAY from the camera. A 2.8mm wide-angle lens covers ~120° horizontal FOV, enough for full cage width from 30+ feet. Chain-link does not interfere with CV — the model learns to ignore the mesh pattern.")

# Architecture diagram
doc.add_paragraph()
add_image_with_caption(doc, img("docmockup_architecture.png"),
    "System architecture: Jetson runs vision AI and game logic. Camera tracks ball. ESP32 drives LED targets. Tablet displays web app.")

doc.add_paragraph()
add_heading_styled(doc, "Game Screens (Actual Software)", level=2)

add_image_with_caption(doc, img("screenshot_attract.png"),
    "Attract Screen — Idle mode with vintage pinball aesthetic. Shows How to Play, scoring rules, multiplier chart, and leaderboards. 'INSERT TOKEN' prompt at center.")

doc.add_paragraph()
add_image_with_caption(doc, img("screenshot_ingame.png"),
    "Game In Progress — Live gameplay showing 5X multiplier, 6,500 score, pitch 6 of 10. Target zones (Home Run 1000, Triple 500, Double 300, Single 200, Bunt 100) on the stadium field. Real-time pitch tracking with ball animation.")

doc.add_page_break()

# ============================================================
# 5. MVP vs FULL BUILD
# ============================================================

add_heading_styled(doc, "5. MVP vs. Full Build", level=1)

add_heading_styled(doc, "MVP (Pilot) — ~$1,600", level=2)
add_body(doc, "Proves the concept with minimal investment and zero disruption to current operations.")

add_styled_table(doc, ["What's Included (MVP)", "What's Not (Phase 1.5+)"], [
    ["Jetson Orin Nano + camera + CV system", "Outdoor TV (use tablet instead)"],
    ["5 LED target panels inside the cage", "Food & beverage integration"],
    ["Web app on mounted tablet outside cage", "Player profiles / mobile app"],
    ["Token machine integration (auto-start)", "Multiple game modes (Blitz only)"],
    ["Daily leaderboard", "Online/cloud leaderboard"],
    ["Pinball sound effects", "QR code score sharing"],
    ["Same token price as regular cages", "Premium pricing"],
], col_widths=[3.2, 3.2])

doc.add_paragraph()

add_heading_styled(doc, "MVP Cost Breakdown", level=2)
add_styled_table(doc, ["Category", "Cost"], [
    ["Compute (Jetson + SSD + PSU + enclosure)", "$389"],
    ["Camera (IMX477 + lens + housing + mount)", "$156"],
    ["LED Targets (ESP32 + LEDs + polycarbonate + frames)", "$414"],
    ["Audio (speakers + USB adapter)", "$100"],
    ["Token Interface", "$40"],
    ["Power & Infrastructure", "$286"],
    ["Mounting Hardware", "$69"],
    ["Tablet display (weatherproof case)", "~$150"],
    ["**MVP TOTAL**", "**~$1,600**"],
], col_widths=[4.5, 1.5])

doc.add_paragraph()
add_callout_box(doc, " Everything in the MVP plus SunBrite Veranda 3 55\" outdoor TV ($999), dedicated surge protector, enhanced audio, and 4G internet connectivity.", bold_prefix="Phase 1.5 (~$3,000):")

doc.add_page_break()

# ============================================================
# 6. BILL OF MATERIALS
# ============================================================

add_heading_styled(doc, "6. Bill of Materials — Key Components", level=1)

add_body(doc, "A detailed 47-line-item BOM with specific product recommendations and Amazon links is available in the companion spreadsheet. Key components:")

add_styled_table(doc, ["#", "Component", "Product", "Price"], [
    ["1.1", "Compute board", "NVIDIA Jetson Orin Nano Super Dev Kit (8GB)", "$249"],
    ["1.3", "Storage", "Samsung 970 EVO Plus 500GB NVMe", "$55"],
    ["2.1", "Camera", "Arducam IMX477 12.3MP CS-Mount (120fps)", "$65"],
    ["2.2", "Lens", "Arducam 2.8mm CS-Mount Wide Angle (~120° FOV)", "$22"],
    ["3.1", "Outdoor TV", "SunBrite Veranda 3 55\" 1000-nit QLED IPX5", "$999"],
    ["4.1", "Target controller", "ESP32-WROOM-32 Dev Board (2-pack)", "$14"],
    ["4.2", "LED strips", "ALITOVE WS2812B IP67 Waterproof (5m) ×6", "$108"],
    ["4.4", "Target covers", "1/4\" Polycarbonate 24\"×24\" ×6", "$132"],
    ["5.1", "Speakers", "Pyle 6.5\" Marine-Grade Waterproof BT Pair", "$85"],
    ["7.1", "Networking", "GL.iNet Spitz 4G LTE Router (OpenWrt)", "$89"],
], col_widths=[0.5, 1.3, 3.2, 0.8])

doc.add_paragraph()

add_styled_table(doc, ["Category", "Subtotal"], [
    ["1 — Compute Unit", "$389"],
    ["2 — Camera System", "$156"],
    ["3 — Spectator Display", "$1,116"],
    ["4 — LED Target System", "$414"],
    ["5 — Audio System", "$100"],
    ["6 — Token Interface", "$40"],
    ["7 — Networking", "$95"],
    ["8 — Power & Infrastructure", "$286"],
    ["9 — Mounting Hardware", "$69"],
    ["**GRAND TOTAL**", "**$2,665**"],
    ["**+ 15% Contingency**", "**~$3,065**"],
], col_widths=[4.0, 1.5])

doc.add_page_break()

# ============================================================
# 7. FINANCIAL PROJECTIONS
# ============================================================

add_heading_styled(doc, "7. Financial Projections", level=1)

add_heading_styled(doc, "MVP Phase — Free Overlay", level=2)
add_body(doc, "During the MVP, the Cage Match cage costs the same tokens as any other cage. Goal: measure engagement, not extract premium pricing.")

add_body(doc, "Key metrics to track:", bold=True)
add_bullet(doc, "Token revenue per hour on Cage Match vs. adjacent standard cages")
add_bullet(doc, "Number of 'replay' sessions (same person buying more tokens immediately)")
add_bullet(doc, "Time spent in spectator area watching")
add_bullet(doc, "Social media mentions / photos of the leaderboard")

add_callout_box(doc, " At ~$1,600 MVP cost and 10 extra token purchases/day at $2 each = 80 days (~3 months).", bold_prefix="Break-even:")

add_heading_styled(doc, "Post-MVP — Premium Pricing", level=2)

add_styled_table(doc, ["Scenario", "Price", "Sessions/Day", "Daily Rev.", "Monthly Rev."], [
    ["Current (standard token)", "$2/token", "~40", "$80", "$2,400"],
    ["MVP (same price, more plays)", "$2/token", "~60 (+50%)", "$120", "$3,600"],
    ["Premium (game session)", "$5/game", "~30", "$150", "$4,500"],
    ["**Group Package (4 players)**", "**$40/group**", "**~8 groups**", "**$320**", "**$9,600**"],
], col_widths=[1.8, 1.0, 1.0, 1.0, 1.2])

doc.add_paragraph()
add_callout_box(doc, " Four friends pay $10 each for 30 minutes of competitive batting — cheaper than TopGolf per person but 4x the revenue of token play per time slot.", bold_prefix="The group package is the real revenue driver.")

add_heading_styled(doc, "Revenue Multiplier Effects", level=2)
add_styled_table(doc, ["Effect", "Impact"], [
    ["**Repeat visits**", "Leaderboard drives 'I need to beat that score' returns"],
    ["**Social media**", "Photos of leaderboard = free marketing"],
    ["**Word of mouth**", "'Have you tried the game cage at Batter Up?'"],
    ["**Extended stays**", "Groups stay longer → more mini golf, snacks, arcade"],
    ["**Birthday packages**", "'Cage Match birthday party' is a differentiated offering"],
    ["**Event bookings**", "Corporate team outings, league nights, tournaments"],
], col_widths=[2.0, 4.5])

doc.add_page_break()

# ============================================================
# 8. IMPLEMENTATION TIMELINE
# ============================================================

add_heading_styled(doc, "8. Implementation Timeline", level=1)
add_body(doc, "Total: ~10 weeks from start to launch", bold=True)

phases = [
    ("Phase 0: Planning & Procurement (Weeks 1–2)", [
        "Order all BOM components",
        "Select cage to convert (recommend: middle-speed, ~50–60 mph)",
        "Survey electrical capacity at target cage",
        "Design target zone layout for specific cage dimensions",
        "Set up Jetson dev environment",
        "Begin CV prototyping (ball detection in test footage)",
    ]),
    ("Phase 1: Hardware Build (Weeks 3–5)", [
        "Fabricate LED target frames (aluminum channel + polycarbonate + LED strips)",
        "Build electronics enclosures (Jetson box + power distribution box)",
        "Wire token machine interface (optocoupler tap)",
        "Mount camera in weatherproof housing",
        "Install ESP32 + LED targets in cage (during off-hours)",
        "Run conduit and power to all components",
        "Mount tablet display outside cage",
    ]),
    ("Phase 2: Software Development (Weeks 3–6, parallel)", [
        "Ball tracking CV pipeline — OpenCV + CUDA on Jetson",
        "Game engine — Python/FastAPI (state machine, scoring, MQTT)",
        "Web app display — React or vanilla JS (game view, leaderboard, animations)",
        "Sound design — pinball hit sounds, multiplier tones, crowd cheers",
        "ESP32 firmware — MQTT listener, LED animations, token detection",
    ]),
    ("Phase 3: Integration & Testing (Weeks 6–7)", [
        "Full system integration test in the cage",
        "Calibrate camera zones to physical target positions",
        "Test in various lighting conditions (day, dusk, night, overcast)",
        "Stress test: continuous play for 2+ hours",
        "Weatherproofing verification",
    ]),
    ("Phase 4: Soft Launch (Week 8)", [
        "Install during off-hours",
        "Family/friends beta testing",
        "Go live — same token price, no fanfare, observe organic engagement",
    ]),
    ("Phase 5: Full Launch (Week 10+)", [
        "Social media announcement",
        "'Grand Opening' event — first weekly tournament",
        "Evaluate premium pricing model based on MVP data",
        "Decide on outdoor TV upgrade",
    ]),
]

for phase_title, items in phases:
    add_heading_styled(doc, phase_title, level=2)
    for item in items:
        add_bullet(doc, item)

doc.add_page_break()

# ============================================================
# 9. COMPETITIVE MOAT
# ============================================================

add_heading_styled(doc, "9. Competitive Moat", level=1)

add_styled_table(doc, ["Moat", "Detail"], [
    ["**Custom hardware**", "Proprietary system built to Batter Up's exact cage dimensions. No off-the-shelf product does this at this price."],
    ["**First-mover on LI**", "Zero gamified cages exist on Long Island. First to market builds brand association."],
    ["**Leaderboard network effect**", "More players = more value for everyone. Players compete against the community, not just themselves."],
    ["**10x cheaper**", "$1,600 MVP vs. HitTrax $10K–$20K vs. ProBatter $30K. Allows experimentation without risk."],
    ["**Software iteration**", "Web app = instant updates. 'Double Points Night' costs nothing to create."],
    ["**Existing foot traffic**", "Batter Up already has customers. This converts them into higher-value visitors."],
], col_widths=[2.0, 4.5])

doc.add_paragraph()

add_heading_styled(doc, "Competitive Position", level=2)
add_body(doc, """
                        Fun / Social
                              |
                   CAGE MATCH |  BATBOX
                   (Batter Up)|  ($7.3M, Dallas only)
                              |
              ────────────────┼────────────────  Technology
                              |
                 Standard     |  HitTrax
                 Batting Cages|  ($10-20K, training)
                              |
""", size=9)
add_body(doc, "Cage Match occupies the 'fun + accessible' quadrant that no one on Long Island currently serves.", italic=True)

# Spectator image
doc.add_paragraph()
add_image_with_caption(doc, img("docmockup_spectators.png"),
    "The social experience: families watch the game unfold, checking scores on the tablet and cheering hits.")

doc.add_page_break()

# ============================================================
# 10. RISKS & MITIGATIONS
# ============================================================

add_heading_styled(doc, "10. Risks & Mitigations", level=1)

add_styled_table(doc, ["Risk", "Severity", "Mitigation"], [
    ["**Ball tracking accuracy**", "Medium", "Test extensively. Night games may be easier than bright sun. Fallback: break-beam sensors."],
    ["**Weather damage**", "Medium", "All components in NEMA 4X enclosures. Camera in IP66 housing. LEDs behind polycarbonate."],
    ["**Ball impact on targets**", "Low", "1/4\" polycarbonate is 250x stronger than glass. LEDs behind panel. Replace covers annually ($22 each)."],
    ["**Customer confusion**", "Low", "'DROP A TOKEN TO PLAY' on idle screen. Targets are intuitive — they light up, you aim."],
    ["**Brother buy-in**", "Low", "MVP is a free overlay. Same tokens, same price. Zero revenue risk."],
    ["**Maintenance**", "Medium", "Jetson auto-boots. Watchdog restarts services. Remote monitoring via 4G. Graceful degradation."],
], col_widths=[1.5, 0.8, 4.2])

doc.add_page_break()

# ============================================================
# 11. PHASE 2 EXPANSION
# ============================================================

add_heading_styled(doc, "11. Phase 2 Expansion", level=1)

add_heading_styled(doc, "Near-Term (3–6 months post-launch)", level=2)
add_bullet(doc, " — replace tablet with SunBrite 55\" display ($1,000)", bold_prefix="Outdoor TV upgrade")
add_bullet(doc, " — QR code scan for persistent stats across visits", bold_prefix="Player profiles")
add_bullet(doc, " — multiplayer head-to-head competition", bold_prefix="Derby mode")
add_bullet(doc, " — Friday Night Derby events, prizes for top scores", bold_prefix="Weekly tournaments")
add_bullet(doc, " — 'Share your score' link after each game", bold_prefix="Social sharing")

add_heading_styled(doc, "Medium-Term (6–12 months)", level=2)
add_bullet(doc, " — different speed for variety", bold_prefix="Second cage conversion")
add_bullet(doc, " — player profiles, leaderboard, session booking", bold_prefix="Mobile app")
add_bullet(doc, " — 'Cage Match Party' as premium offering", bold_prefix="Birthday party packages")
add_bullet(doc, " — triangulated 3D tracking for exit velocity + launch angle", bold_prefix="Second camera")
add_bullet(doc, " — display pitch speed and exit velocity", bold_prefix="Speed radar integration")

add_heading_styled(doc, "Long-Term (12–24 months)", level=2)
add_bullet(doc, " — food truck or concession stand near game cages", bold_prefix="Food & beverage")
add_bullet(doc, " — team-building packages", bold_prefix="Corporate events")
add_bullet(doc, " — weekly competitive leagues with standings", bold_prefix="League nights")
add_bullet(doc, " — 2–4 game cages competing simultaneously", bold_prefix="Multi-cage tournament mode")
add_bullet(doc, " — sell/license Cage Match system to other operators", bold_prefix="Franchise the system")

doc.add_paragraph()
add_callout_box(doc, " Batter Up becomes the first 'sports entertainment' batting facility on Long Island — a destination, not a commodity. The mini golf draws families; Cage Match draws young adults, groups, and repeat visitors. A food truck adds dwell time. Weekly tournaments create a community. This is the Bowlero playbook applied to batting cages.", bold_prefix="The Endgame Vision:")

doc.add_page_break()

# ============================================================
# 12. VISUAL MOCKUPS
# ============================================================

add_heading_styled(doc, "12. Visual Mockups", level=1)

add_body(doc, "The following AI-generated concept renderings visualize the Cage Match experience at Batter Up.")

doc.add_paragraph()

add_heading_styled(doc, "Cage Exterior — The Game Cage vs. Standard Cages", level=2)
add_image_with_caption(doc, img("docmockup_exterior.png"),
    "The Cage Match cage stands out from adjacent standard cages. LED targets glow inside. A tablet display shows scores to spectators. Red token machine at entrance — same as every other cage.")

add_heading_styled(doc, "Batter's Point of View — Inside the Cage", level=2)
add_image_with_caption(doc, img("docmockup_batter_pov.png"),
    "What the batter sees: 5 illuminated target panels with point values. Gold bullseye center (1,000), red (500), blue (300), green (200), yellow (100). The layout rewards both power and placement.")

doc.add_page_break()

add_heading_styled(doc, "Tablet Display — Game in Progress", level=2)
add_image_with_caption(doc, img("docmockup_tablet_ui.png"),
    "The Cage Match web app: player score with multiplier, pitch counter, and daily leaderboard. At idle: 'DROP A TOKEN TO PLAY.'", width=Inches(3.5))

add_heading_styled(doc, "System Architecture", level=2)
add_image_with_caption(doc, img("docmockup_architecture.png"),
    "Hardware architecture: Jetson Orin Nano runs vision AI and game logic. Camera tracks ball. ESP32 drives LED targets via MQTT. Tablet connects via WiFi.")

doc.add_page_break()

add_heading_styled(doc, "The Spectator Experience", level=2)
add_image_with_caption(doc, img("docmockup_spectators.png"),
    "The TopGolf effect: families and friends watch the game, check scores on the tablet, film with phones. Spectators are part of the experience.")

# Actual facility photos
doc.add_paragraph()
add_heading_styled(doc, "Batter Up — Current Facility", level=2)

add_image_with_caption(doc, img("../images/bu30.jpeg"),
    "Inside the cages today — kids batting with red and blue shade sails overhead. Standard chain-link, no technology.", width=Inches(5.5))

add_image_with_caption(doc, img("../images/bu31.jpeg"),
    "Wide view showing blue shade sails, red tent structures, and the snack bar / mini golf area adjacent to the cages.", width=Inches(5.5))

# ============================================================
# FOOTER / CLOSING
# ============================================================

doc.add_page_break()

doc.add_paragraph()
doc.add_paragraph()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run("━" * 40)
run.font.color.rgb = RGBColor(0xa0, 0xc4, 0xff)
run.font.size = Pt(14)

doc.add_paragraph()

closing2 = doc.add_paragraph()
closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing2.add_run("CAGE MATCH")
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3e)

closing3 = doc.add_paragraph()
closing3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing3.add_run("Prepared by Nick DeMarco  •  April 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

closing4 = doc.add_paragraph()
closing4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing4.add_run("Batter Up — Bethpage, Long Island")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ============================================================
# SAVE
# ============================================================

output_path = os.path.join(BASE, "Cage_Match_Business_Plan.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
